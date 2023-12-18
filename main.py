import sys
import random
import pandas as pd
from docx import Document
from docx.shared import Pt
from PyQt6.QtWidgets import QApplication, QMainWindow, QLabel, QPushButton, QVBoxLayout, QWidget, QFormLayout, \
    QComboBox, QTimeEdit, QMessageBox, QTextEdit, QTableWidget
from PyQt6.QtCore import QTime, Qt
import qrcode
from PyQt6.QtGui import QPixmap
from PIL import Image
import sqlite3
class TrainScheduleApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Расписание электричек")

        self.setGeometry(100, 100, 400, 200)


        # Создаем главный виджет и вертикальное размещение
        central_widget = QWidget(self)
        layout = QVBoxLayout(central_widget)

        # Добавляем заголовок
        title_label = QLabel("Билет на электричку")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        title_label.setStyleSheet("font-size: 18px; font-weight: bold;color: #8B0000")
        layout.addWidget(title_label)

        # Создаем форму для выбора параметров
        image = Image.open("train.jpg")

        # Уменьшаем размер изображения
        new_size = (500, 300)  # Задайте новые размеры
        resized_image = image.resize(new_size)

        # Сохраняем уменьшенное изображение
        resized_image.save("train_resized.jpg")
        form_layout = QFormLayout()
        train_image = QLabel(self)
        pixmap = QPixmap("train_resized.jpg")
        train_image.setPixmap(pixmap)
        train_image.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(train_image)

        self.setCentralWidget(central_widget)
        self.departure_combo = QComboBox()
        self.departure_combo.addItems(["Выхино", "Воронок", "Рязань", "Тула", "Автозаводская", "Панки"])
        form_layout.addRow("Откуда:", self.departure_combo)

        self.arrival_combo = QComboBox()
        self.arrival_combo.addItems(["Выхино", "Воронок", "Рязань", "Тула", "Автозаводская", "Панки"])
        form_layout.addRow("Куда:", self.arrival_combo)

        self.departure_time_edit = QTimeEdit()
        self.departure_time_edit.setDisplayFormat("HH:mm")
        self.departure_time_edit.setTime(QTime.currentTime())
        form_layout.addRow("Время отправления:", self.departure_time_edit)

        self.arrival_time_label = QLabel()
        form_layout.addRow("Время прибытия:", self.arrival_time_label)

        save_button = QPushButton("Сохранить")
        save_button.clicked.connect(self.save_ticket)
        form_layout.addRow("", save_button)


        layout.addLayout(form_layout)
        layout.addStretch()

        self.database_content = QTextEdit()
        layout.addWidget(self.database_content)
        self.setCentralWidget(central_widget)
        self.get_database_content()

    def get_database_content(self):

        conn = sqlite3.connect('train_schedule.db')
        cursor = conn.cursor()

        cursor.execute('''
               SELECT * FROM ticket
           ''')

        data = cursor.fetchall()

        table = QTableWidget()

        columns = ["№", "Train Number", "Departure", "Arrival", "Departure Time", "Arrival Time", "Cost"]

        table.setHorizontalHeaderLabels(columns)

        df = pd.DataFrame(data, columns=columns)

        self.database_content.setPlainText(str(df))

        conn.close()

    def save_ticket(self):
        departure = self.departure_combo.currentText()
        arrival = self.arrival_combo.currentText()

        if departure == arrival:
            QMessageBox.critical(self, "Ошибка", "Выберите разные станции.")
            # Установка таймера на 3 секунды


            return

        else:
            self.departure_combo.setStyleSheet("")
            self.arrival_combo.setStyleSheet("")

        departure_time = self.departure_time_edit.time()
        departure_time_str = departure_time.toString("HH:mm")


        # Расчет стоимости
        cost = 0
        if departure == "Выхино" and arrival == "Панки":
            cost = 25
        elif departure == "Выхино" and arrival == "Рязань":
            cost = 30
        elif departure == "Выхино" and arrival == "Воронок":
            cost = 35
        elif departure == "Выхино" and arrival == "Автозаводская":
            cost = 30
        elif departure == "Выхино" and arrival == "Тула":
            cost = 40
        elif departure == "Воронок" and arrival == "Рязань":
            cost = 70
        elif departure == "Воронок" and arrival == "Выхино":
            cost = 35
        elif departure == "Воронок" and arrival == "Панки":
            cost = 50
        elif departure == "Воронок" and arrival == "Автозаводская":
            cost = 45
        elif departure == "Воронок" and arrival == "Тула":
            cost = 55
        elif departure == "Рязань" and arrival == "Автозаводская":
            cost = 40
        elif departure == "Рязань" and arrival == "Выхино":
            cost = 30
        elif departure == "Рязань" and arrival == "Панки":
            cost = 50
        elif departure == "Рязань" and arrival == "Тула":
            cost = 25
        elif departure == "Рязань" and arrival == "Воронок":
            cost = 70
        elif departure == "Тула" and arrival == "Выхино":
            cost = 40
        elif departure == "Тула" and arrival == "Рязань":
            cost = 25
        elif departure == "Тула" and arrival == "Панки":
            cost = 40
        elif departure == "Тула" and arrival == "Автозаводская":
            cost = 45
        elif departure == "Тула" and arrival == "Воронок":
            cost = 55
        elif departure == "Автозаводская" and arrival == "Панки":
            cost = 25
        elif departure == "Автозаводская" and arrival == "Тула":
            cost = 45
        elif departure == "Автозаводская" and arrival == "Рязань":
            cost = 40
        elif departure == "Автозаводская" and arrival == "Выхино":
            cost = 30
        elif departure == "Автозаводская" and arrival == "Воронок":
            cost = 45
        elif departure == "Панки" and arrival == "Воронок":
            cost = 50
        elif departure == "Панки" and arrival == "Выхино":
            cost = 25
        elif departure == "Панки" and arrival == "Рязань":
            cost = 50
        elif departure == "Панки" and arrival == "Тула":
            cost = 45
        elif departure == "Панки" and arrival == "Автозаводская":
            cost = 25

        train_number = random.randint(1000, 9999)

        # Расчет и установка времени прибытия
        arrival_time = departure_time.addSecs(cost * 60)
        arrival_time_str = arrival_time.toString("HH:mm")
        self.arrival_time_label.setText(arrival_time_str)

        ticket_info = f" Номер поезда: {train_number}\n   "

        ticket_info += f"Маршрут: {departure} - {arrival}\n   "

        ticket_info += f"Время отправления: {departure_time_str}\n   "

        ticket_info += f"Время прибытия: {arrival_time_str}\n   "
        ticket_info += f"Стоимость: {cost}"
        doc = Document()
        qr_data = f"Номер поезда: {train_number}\n Отправка: {departure}\nПрибытие: {arrival}\nВремя отправки: {departure_time_str}\nВремя прибытия: {arrival_time_str}\nЦена: {cost}"
        qr_code = qrcode.make(qr_data)
        qr_code.save("qrcode.png")
        doc.add_picture("qrcode.png", width=Pt(150), height=Pt(150))

        doc.add_paragraph(ticket_info)
        doc.save("ticket.docx")

        with open("ticket.txt", "w", encoding="utf-8") as ticket_file:
            ticket_file.write(ticket_info)

        df = pd.DataFrame({'Расписание электрички вместе с билетом': [ticket_info]})
        df.to_excel("ticket.xlsx", index=False)

        conn = sqlite3.connect('train_schedule.db')
        cursor = conn.cursor()



        cursor.execute('''
                    CREATE TABLE IF NOT EXISTS ticket (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        train_number INTEGER,
                        departure TEXT,
                        arrival TEXT,
                        departure_time TEXT,
                        arrival_time TEXT,
                        cost INTEGER
                    )
                ''')

        # Вставка записи в таблицу "tickets"
        cursor.execute('''
                    INSERT INTO ticket (train_number, departure, arrival, departure_time, arrival_time, cost)
                    VALUES (?, ?, ?, ?, ?, ?)
                ''', (train_number, departure, arrival, departure_time_str, arrival_time_str, cost))

        # Сохранение изменений и закрытия соединения

        conn.commit()
        conn.close()


    print("Билет сохранен!")


if __name__ == "__main__":
    app = QApplication(sys.argv)

    # Задаем стиль для кнопок
    app.setStyleSheet(
        "QPushButton {"
        "font-size: 14px;"
        "padding: 8px 16px;"
        "border-radius: 4px;"
        "background-color: #FF6680;"
        "color: white;"
        "}"
        "QPushButton:hover {"
        "background-color: #FF3355;"
        "}"
        "QComboBox {"
        "font-size: 14px;"
        "padding: 8px 16px;"
        "border-radius: 4px;"
        "}"
    )

    schedule_app = TrainScheduleApp()
    schedule_app.show()
    sys.exit(app.exec())